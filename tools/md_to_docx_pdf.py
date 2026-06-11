"""Minimal Markdown -> .docx (+ .pdf) converter for battle-doc deliverables.
Handles: # / ## headings, | tables |, ``` code blocks, - bullets, 1. numbered,
**bold** inline, *italic* inline, horizontal rule (---), plain paragraphs.
"""
import re, sys, os

# ---- fail-closed factual gate (primary enforcement) ----
# Re-lint the EXACT file we are about to convert, before any conversion work.
# This closes the lint-then-modify-then-export TOCTOU window: the file we open
# here is the file we check. A HARD FAIL aborts the export (exit 1). See
# tools/lint_team.py. Conversion is the single canonical export path for
# deliverables; do not add alternate exporters.
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import lint_team as _lint  # noqa: E402

_SRC = sys.argv[1] if len(sys.argv) > 1 else None
if _SRC and os.path.exists(_SRC):
    _res = _lint.lint_for_export(_SRC)
    if not _res.ok:
        sys.stderr.write(
            "\nEXPORT BLOCKED by tools/lint_team.py - fix these before converting:\n\n"
            + _res.report() + "\n\n"
            "(This is the factual gate. If a flag is a confirmed data gap, add it to "
            "data/lint_allowlists.json with a reason. Tera/dash/customer-language flags "
            "must be fixed in the draft.)\n")
        _lint._log_denial(_SRC, _res)
        sys.exit(1)
    elif _res.warn:
        sys.stderr.write("lint warnings (not blocking):\n" + _res.report() + "\n\n")

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
                                TableStyle, ListFlowable, ListItem, Preformatted, HRFlowable)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT

SRC, DOCX, PDF = sys.argv[1], sys.argv[2], sys.argv[3]

lines = open(SRC, encoding="utf-8").read().split("\n")

# ---- parse into blocks ----
blocks = []  # (type, payload)
i = 0
while i < len(lines):
    ln = lines[i]
    if ln.strip().startswith("```"):
        buf = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            buf.append(lines[i]); i += 1
        i += 1
        blocks.append(("code", buf)); continue
    if ln.strip().startswith("|") and "|" in ln:
        buf = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            buf.append(lines[i]); i += 1
        rows = []
        for r in buf:
            cells = [c.strip() for c in r.strip().strip("|").split("|")]
            if all(set(c) <= set("-: ") for c in cells):  # separator row
                continue
            rows.append(cells)
        blocks.append(("table", rows)); continue
    if re.match(r"^#{1,6}\s", ln):
        lvl = len(ln) - len(ln.lstrip("#"))
        blocks.append(("h%d" % lvl, ln.lstrip("#").strip())); i += 1; continue
    if ln.strip() in ("---", "***", "___"):
        blocks.append(("hr", None)); i += 1; continue
    if re.match(r"^\s*[-*]\s+", ln):
        buf = []
        while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
            buf.append(re.sub(r"^\s*[-*]\s+", "", lines[i])); i += 1
        blocks.append(("ul", buf)); continue
    if re.match(r"^\s*\d+\.\s+", lines[i]):
        buf = []
        while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
            buf.append(re.sub(r"^\s*\d+\.\s+", "", lines[i])); i += 1
        blocks.append(("ol", buf)); continue
    if ln.strip() == "":
        i += 1; continue
    # paragraph (gather consecutive non-empty, non-special lines)
    buf = [ln]
    i += 1
    while i < len(lines) and lines[i].strip() and not re.match(r"^(\s*[-*]\s+|\s*\d+\.\s+|#|\||```)", lines[i]) and lines[i].strip() not in ("---","***","___"):
        buf.append(lines[i]); i += 1
    blocks.append(("p", " ".join(buf)))

def split_runs(text):
    """Yield (text, bold, italic) tuples from **bold**/*italic* markup."""
    parts = re.split(r"(\*\*.+?\*\*|\*.+?\*)", text)
    for p in parts:
        if not p:
            continue
        if p.startswith("**") and p.endswith("**"):
            yield (p[2:-2], True, False)
        elif p.startswith("*") and p.endswith("*"):
            yield (p[1:-1], False, True)
        else:
            yield (p, False, False)

# ================= DOCX =================
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(11)
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Inches(0.8)
    sec.left_margin = sec.right_margin = Inches(0.9)

def add_runs(p, text):
    for t, b, it in split_runs(text):
        r = p.add_run(t); r.bold = b; r.italic = it

for kind, payload in blocks:
    if kind == "h1":
        h = doc.add_heading(level=0); add_runs(h, payload)
    elif kind == "h2":
        h = doc.add_heading(level=1); add_runs(h, payload)
    elif kind == "h3":
        h = doc.add_heading(level=2); add_runs(h, payload)
    elif kind == "p":
        add_runs(doc.add_paragraph(), payload)
    elif kind == "ul":
        for it in payload:
            add_runs(doc.add_paragraph(style="List Bullet"), it)
    elif kind == "ol":
        for it in payload:
            add_runs(doc.add_paragraph(style="List Number"), it)
    elif kind == "code":
        p = doc.add_paragraph()
        r = p.add_run("\n".join(payload)); r.font.name = "Consolas"; r.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.2)
    elif kind == "hr":
        doc.add_paragraph("_" * 60)
    elif kind == "table":
        rows = payload
        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
        t.style = "Light Grid Accent 1"
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                c = t.cell(ri, ci); c.paragraphs[0].clear()
                add_runs(c.paragraphs[0], cell)
                if ri == 0:
                    for rn in c.paragraphs[0].runs:
                        rn.bold = True
doc.save(DOCX)

# ================= PDF =================
ss = getSampleStyleSheet()
body = ParagraphStyle("body", parent=ss["Normal"], fontName="Helvetica", fontSize=10, leading=14, spaceAfter=6)
h1 = ParagraphStyle("h1", parent=ss["Title"], fontName="Helvetica-Bold", fontSize=20, spaceAfter=12)
h2 = ParagraphStyle("h2", parent=ss["Heading2"], fontName="Helvetica-Bold", fontSize=14, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1F4E79"))
code = ParagraphStyle("code", parent=ss["Code"], fontName="Courier", fontSize=8.5, leading=10.5, leftIndent=10, backColor=colors.HexColor("#F2F2F2"))

def md_inline(text):
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", text)
    return text

flow = []
for kind, payload in blocks:
    if kind == "h1":
        flow.append(Paragraph(md_inline(payload), h1))
    elif kind in ("h2", "h3"):
        flow.append(Paragraph(md_inline(payload), h2))
    elif kind == "p":
        flow.append(Paragraph(md_inline(payload), body))
    elif kind == "ul":
        flow.append(ListFlowable([ListItem(Paragraph(md_inline(x), body)) for x in payload], bulletType="bullet", start="•"))
    elif kind == "ol":
        flow.append(ListFlowable([ListItem(Paragraph(md_inline(x), body)) for x in payload], bulletType="1"))
    elif kind == "code":
        flow.append(Preformatted("\n".join(payload), code))
        flow.append(Spacer(1, 6))
    elif kind == "hr":
        flow.append(Spacer(1, 4)); flow.append(HRFlowable(width="100%", color=colors.grey)); flow.append(Spacer(1, 4))
    elif kind == "table":
        rows = payload
        data = [[Paragraph(md_inline(c), ParagraphStyle("c", parent=body, fontSize=8.5, leading=11,
                 fontName=("Helvetica-Bold" if ri == 0 else "Helvetica"))) for c in row]
                for ri, row in enumerate(rows)]
        ncols = len(rows[0])
        avail = letter[0] - 1.8*inch
        # weight first and Moves columns wider for the 5-col team table
        if ncols == 5:
            widths = [w/100.0*avail for w in (15, 13, 14, 38, 20)]
        else:
            widths = [avail/ncols]*ncols
        t = RLTable(data, colWidths=widths, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1F4E79")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#AAAAAA")),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#EDF2F8")]),
            ("TOPPADDING", (0,0), (-1,-1), 3), ("BOTTOMPADDING", (0,0), (-1,-1), 3),
            ("LEFTPADDING", (0,0), (-1,-1), 5), ("RIGHTPADDING", (0,0), (-1,-1), 5),
        ]))
        flow.append(t); flow.append(Spacer(1, 8))

SimpleDocTemplate(PDF, pagesize=letter, topMargin=0.7*inch, bottomMargin=0.7*inch,
                  leftMargin=0.8*inch, rightMargin=0.8*inch).build(flow)
print("wrote", DOCX, "and", PDF)
